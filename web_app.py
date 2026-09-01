import json
import re
import secrets
from calendar import monthrange
from io import BytesIO
from datetime import date, datetime, timedelta
from pathlib import Path
from urllib.parse import quote, urlencode

from flask import (
    Flask,
    abort,
    g,
    has_request_context,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)
from config import PROJECT_VERSION
from utils.auth import environment_value, load_secret_key
from utils.article_reader import extract_article, yahoo_article_is_polluted
from utils.diagnostics import alert_summary, source_alerts, system_alerts
from utils.keywords import (
    add_keyword,
    load_keywords,
    rebuild_found_news,
    remove_keyword,
)
from utils.logger import error_log_stats, get_logger, read_recent_errors
from utils.proxy import kyodo_proxy_status
from utils.security import AttemptLimiter
from utils.source_groups import (
    AGENCIES_GROUP,
    AGENCY_SOURCES,
    GOVERNMENT_GROUP,
    GOVERNMENT_SOURCES,
    NEWSPAPERS_GROUP,
    NEWSPAPER_SOURCES,
    is_yahoo_source,
    source_group as get_source_group,
)
from utils.storage import (
    authenticate_user,
    create_manual_backup,
    create_dictionary_deck,
    bookmarked_urls,
    count_users,
    count_bookmarks,
    create_bookmark_folder,
    delete_collection_note,
    delete_calendar_event,
    delete_personal_note,
    create_user,
    delete_user,
    delete_bookmark_folder,
    enqueue_parser_job,
    ensure_favorites_folder,
    find_news_by_url,
    list_users,
    list_database_backups,
    list_source_incidents,
    list_parser_jobs,
    list_bookmark_folders,
    list_bookmarks,
    list_collection_bookmarks,
    list_collection_note_read_ids,
    list_collection_notes,
    list_calendar_events,
    list_dictionary_cards,
    list_dictionary_decks,
    list_personal_notes,
    list_news_index,
    list_unread_news_index,
    list_news_page,
    list_shared_collections,
    load_collection,
    load_source_order,
    load_source_settings,
    load_all_news,
    load_cached_article,
    load_found_news,
    load_user,
    migrate_legacy_unread,
    remove_bookmark,
    rename_bookmark_folder,
    save_cached_article,
    save_bookmark,
    save_bookmark_folder_order,
    save_collection_note,
    save_calendar_event,
    save_dictionary_card,
    save_personal_note,
    save_external_bookmark,
    save_source_order,
    set_source_enabled,
    set_collection_note_read,
    mark_news_group_read,
    mark_news_read,
    set_user_active,
    set_user_password,
    set_user_role,
    review_dictionary_card,
    update_collection,
    update_collection_note,
    news_group_counts,
    news_unread_summary,
    news_source_counts,
    source_news_statistics,
    source_incident_statistics,
    source_reliability_statistics,
    database_stats,
    update_bookmark,
)


app = Flask(__name__)
PROJECT_DIR = Path(__file__).resolve().parent
NEWS_PER_PAGE = 20
UNREAD_INDEX_LIMIT = 2000
FAST_NAVIGATION_ENDPOINTS = {
    "index",
    "found_page",
    "filter_source",
    "agencies_page",
    "agencies_found_page",
    "agencies_filter_source",
    "newspapers_page",
    "newspapers_found_page",
    "newspapers_filter_source",
    "article_page",
    "bookmarks_page",
}


def _enabled_setting(name):
    return environment_value(name).casefold() in {"1", "true", "yes", "on"}


def _server_host():
    """Возвращает интерфейс веб-сервера; по умолчанию доступ только локальный."""
    return environment_value("MONITOR_HOST", "127.0.0.1").strip() or "127.0.0.1"


def _allowed_hosts():
    configured = environment_value("MONITOR_ALLOWED_HOSTS")
    return {
        host.strip().casefold().rstrip(".")
        for host in configured.split(",")
        if host.strip()
    }


app.config.update(
    SECRET_KEY=load_secret_key(),
    PERMANENT_SESSION_LIFETIME=timedelta(days=14),
    MAX_CONTENT_LENGTH=1_000_000,
    SESSION_COOKIE_NAME="news_monitor_session",
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=_enabled_setting("MONITOR_HTTPS"),
    AUTH_DISABLED=_enabled_setting("MONITOR_AUTH_DISABLED"),
    ALLOWED_HOSTS=_allowed_hosts(),
)
security_logger = get_logger("web_security")
LOGIN_PAIR_LIMITER = AttemptLimiter(5, 15 * 60, 15 * 60)
LOGIN_ACCOUNT_LIMITER = AttemptLimiter(10, 15 * 60, 30 * 60)
LOGIN_IP_LIMITER = AttemptLimiter(20, 15 * 60, 30 * 60)


def load_json(filename, default):
    # Новости теперь живут в SQLite. Имя функции оставлено прежним,
    # чтобы маршруты и тесты интерфейса не пришлось переписывать целиком.
    collection_loaders = {
        "all_news.json": load_all_news,
        "found_news.json": load_found_news,
    }
    loader = collection_loaders.get(filename)
    if loader is not None:
        request_cache = None
        if has_request_context():
            request_cache = getattr(g, "_news_collection_cache", None)
            if request_cache is None:
                request_cache = {}
                g._news_collection_cache = request_cache
            if filename in request_cache:
                return request_cache[filename]
        result = loader()
        if request_cache is not None:
            request_cache[filename] = result
        return result

    path = PROJECT_DIR / filename
    if not path.exists():
        return default
    try:
        with path.open("r", encoding="utf-8") as file:
            data = json.load(file)
        return data
    except (OSError, json.JSONDecodeError):
        return default


def csrf_token():
    """Возвращает CSRF-токен текущей сессии, создавая его при необходимости."""
    token = session.get("_csrf_token")
    if not token:
        token = secrets.token_urlsafe(32)
        session["_csrf_token"] = token
    return token


def csrf_is_valid():
    """Проверяет токен для изменяющих состояние запросов."""
    if app.config.get("AUTH_DISABLED"):
        return True
    expected = str(session.get("_csrf_token", ""))
    supplied = str(
        request.form.get("csrf_token", "")
        or request.headers.get("X-CSRF-Token", "")
    )
    return bool(expected and supplied and secrets.compare_digest(expected, supplied))


def current_user():
    """Возвращает вошедшего пользователя для проверок прав."""
    return getattr(g, "current_user", None)


def safe_next_url(value):
    """Разрешает возврат после входа только на локальный URL приложения."""
    value = str(value or "").strip()
    if value.startswith("/") and not value.startswith("//"):
        return value
    return "/"


def _request_host():
    """Возвращает имя узла без порта для строгой проверки Host."""
    host = str(request.host or "").strip().casefold().rstrip(".")
    if host.startswith("["):
        return host.partition("]")[0] + "]"
    return host.rsplit(":", 1)[0] if ":" in host else host


def _client_address():
    return str(request.remote_addr or "unknown")[:80]


def _login_limit_keys(username):
    address = _client_address()
    account = " ".join(str(username or "").split()).casefold()[:80]
    return (
        (LOGIN_PAIR_LIMITER, f"{address}|{account}"),
        (LOGIN_ACCOUNT_LIMITER, account),
        (LOGIN_IP_LIMITER, address),
    )


def _login_retry_after(username):
    return max(
        limiter.retry_after(key)
        for limiter, key in _login_limit_keys(username)
    )


def _register_login_failure(username):
    return max(
        limiter.register_failure(key)
        for limiter, key in _login_limit_keys(username)
    )


def _clear_login_failures(username):
    for limiter, key in _login_limit_keys(username):
        limiter.clear(key)


@app.before_request
def validate_request_host():
    """Отклоняет подменённый Host после включения публичного режима."""
    allowed = app.config.get("ALLOWED_HOSTS") or set()
    if allowed and _request_host() not in allowed | {"127.0.0.1", "localhost"}:
        abort(400)


@app.before_request
def load_current_user():
    """Загружает сессию и закрывает приложение от анонимного доступа."""
    if app.config.get("AUTH_DISABLED"):
        g.current_user = {
            "id": 0,
            "username": "Тест",
            "role": "admin",
            "is_active": True,
        }
        return None

    user = load_user(session.get("user_id"))
    if user and user.get("is_active"):
        g.current_user = user
    else:
        session.pop("user_id", None)
        g.current_user = None

    endpoint = request.endpoint or ""
    if endpoint in {"static", "healthz"}:
        return None

    users_exist = count_users() > 0
    if not users_exist:
        if app.config.get("ALLOWED_HOSTS") and not _enabled_setting(
            "MONITOR_ALLOW_REMOTE_SETUP"
        ):
            abort(403)
        if endpoint != "setup":
            return redirect(url_for("setup"))
        return None

    if endpoint == "setup":
        return redirect(url_for("index") if g.current_user else url_for("login"))
    if endpoint == "login":
        if g.current_user:
            return redirect(url_for("index"))
        return None
    if not g.current_user:
        return redirect(
            url_for("login", next=safe_next_url(request.full_path or request.path))
        )
    return None


@app.after_request
def add_security_headers(response):
    """Добавляет безопасные браузерные настройки ко всем ответам."""
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "DENY")
    response.headers.setdefault("Referrer-Policy", "same-origin")
    response.headers.setdefault(
        "Permissions-Policy",
        "camera=(), microphone=(), geolocation=(), payment=()",
    )
    response.headers.setdefault(
        "Content-Security-Policy",
        "default-src 'self'; "
        "base-uri 'self'; "
        "connect-src 'self'; "
        "form-action 'self'; "
        "frame-ancestors 'none'; "
        "img-src 'self' data:; "
        "object-src 'none'; "
        "script-src 'self' 'unsafe-inline'; "
        "style-src 'self' 'unsafe-inline'",
    )
    if request.endpoint != "static":
        fast_navigation_response = (
            request.method == "GET"
            and request.endpoint in FAST_NAVIGATION_ENDPOINTS
            and response.status_code == 200
        )
        if fast_navigation_response:
            # Это только приватный кэш конкретного браузера. Короткое окно
            # убирает повторную загрузку после статьи и позволяет безопасно
            # подогреть соседний раздел, не кэшируя вход и админские формы.
            response.headers.setdefault(
                "Cache-Control",
                "private, max-age=30, must-revalidate",
            )
            response.vary.add("Cookie")
        else:
            response.headers.setdefault("Cache-Control", "no-store")
    if app.config.get("SESSION_COOKIE_SECURE") and request.is_secure:
        response.headers.setdefault(
            "Strict-Transport-Security", "max-age=31536000; includeSubDomains"
        )
    return response


@app.get("/healthz")
def healthz():
    """Лёгкая проверка для Nginx и systemd без раскрытия данных проекта."""
    return jsonify(status="ok")


@app.route("/setup", methods=["GET", "POST"])
def setup():
    """На первом запуске создаёт единственного первого администратора."""
    error = ""
    username = str(request.form.get("username", "")).strip()
    if request.method == "POST":
        if not csrf_is_valid():
            abort(400)
        password = request.form.get("password", "")
        password_confirm = request.form.get("password_confirm", "")
        if password != password_confirm:
            error = "Пароли не совпадают"
        else:
            try:
                user = create_user(username, password, role="admin")
            except ValueError as creation_error:
                error = str(creation_error)
            else:
                session.clear()
                session.permanent = True
                session["user_id"] = user["id"]
                csrf_token()
                return redirect(url_for("index"))
    return render_template(
        "auth.html",
        mode="setup",
        error=error,
        username=username,
        csrf_token=csrf_token(),
        next_url="/",
    )


@app.route("/login", methods=["GET", "POST"])
def login():
    """Открывает защищённую сессию после проверки логина и пароля."""
    error = ""
    retry_after = 0
    username = str(request.form.get("username", "")).strip()
    next_url = safe_next_url(request.values.get("next", "/"))
    if request.method == "POST":
        if not csrf_is_valid():
            abort(400)
        retry_after = _login_retry_after(username)
        if retry_after:
            error = "Слишком много попыток. Попробуйте немного позже."
        else:
            user = authenticate_user(username, request.form.get("password", ""))
            if user is None:
                retry_after = _register_login_failure(username)
                error = (
                    "Слишком много попыток. Попробуйте немного позже."
                    if retry_after
                    else "Неверный логин или пароль"
                )
                security_logger.info(
                    "Неудачный вход: пользователь=%r адрес=%s",
                    " ".join(username.split())[:80],
                    _client_address(),
                )
                if retry_after:
                    security_logger.warning(
                        "Временная блокировка входа: пользователь=%r адрес=%s",
                        " ".join(username.split())[:80],
                        _client_address(),
                    )
            else:
                _clear_login_failures(username)
                security_logger.info(
                    "Успешный вход: пользователь=%r адрес=%s",
                    user["username"],
                    _client_address(),
                )
                session.clear()
                session.permanent = True
                session["user_id"] = user["id"]
                csrf_token()
                return redirect(next_url)
    rendered = render_template(
        "auth.html",
        mode="login",
        error=error,
        username=username,
        csrf_token=csrf_token(),
        next_url=next_url,
    )
    if retry_after:
        return rendered, 429, {"Retry-After": str(retry_after)}
    return rendered


@app.post("/logout")
def logout():
    """Завершает текущую пользовательскую сессию."""
    if not csrf_is_valid():
        abort(400)
    session.clear()
    response = redirect(url_for("login"))
    # Приватный кэш ускоряет навигацию только внутри активной сессии.
    # При выходе браузеру явно поручается удалить сохранённые страницы.
    response.headers["Clear-Site-Data"] = '"cache"'
    return response


@app.route("/account", methods=["GET", "POST"])
def account():
    """Позволяет вошедшему пользователю безопасно сменить свой пароль."""
    user = current_user()
    error = ""
    message = str(request.args.get("message", "")).strip()
    if request.method == "POST":
        if not csrf_is_valid():
            abort(400)
        current_password = request.form.get("current_password", "")
        new_password = request.form.get("new_password", "")
        password_confirm = request.form.get("password_confirm", "")
        if authenticate_user(user["username"], current_password) is None:
            error = "Текущий пароль введён неверно"
        elif new_password != password_confirm:
            error = "Новые пароли не совпадают"
        else:
            try:
                set_user_password(user["id"], new_password)
            except ValueError as password_error:
                error = str(password_error)
            else:
                session.clear()
                session.permanent = True
                session["user_id"] = user["id"]
                csrf_token()
                return redirect(
                    url_for("account", message="Пароль успешно изменён")
                )
    return render_template(
        "settings.html",
        mode="account",
        title="Мой аккаунт",
        subtitle=f"{user['username']} · "
        + ("администратор" if user["role"] == "admin" else "пользователь"),
        current_user=user,
        csrf_token=csrf_token(),
        message=message,
        error=error,
    )


@app.route("/admin/users", methods=["GET", "POST"])
def admin_users():
    """Управляет аккаунтами и позволяет администратору удалить чужой аккаунт."""
    administrator = current_user()
    if not administrator or administrator.get("role") != "admin":
        abort(403)

    if request.method == "POST":
        if not csrf_is_valid():
            abort(400)
        action = str(request.form.get("action", "")).strip()
        try:
            if action == "create":
                password = request.form.get("password", "")
                if password != request.form.get("password_confirm", ""):
                    raise ValueError("Пароли нового пользователя не совпадают")
                created = create_user(
                    request.form.get("username", ""),
                    password,
                    request.form.get("role", "user"),
                )
                message = f"Пользователь {created['username']} создан"
            else:
                target = load_user(request.form.get("user_id"))
                if target is None:
                    raise ValueError("Пользователь не найден")
                if action == "role":
                    if target["id"] == administrator["id"]:
                        raise ValueError("Нельзя менять собственную роль")
                    updated = set_user_role(target["id"], request.form.get("role"))
                    message = f"Роль пользователя {updated['username']} изменена"
                elif action == "toggle":
                    if target["id"] == administrator["id"]:
                        raise ValueError("Нельзя отключить собственный аккаунт")
                    updated = set_user_active(target["id"], not target["is_active"])
                    state = "включён" if updated["is_active"] else "отключён"
                    message = f"Вход для {updated['username']} {state}"
                elif action == "password":
                    updated = set_user_password(
                        target["id"], request.form.get("password", "")
                    )
                    message = f"Пароль пользователя {updated['username']} изменён"
                elif action == "delete":
                    if target["id"] == administrator["id"]:
                        raise ValueError("Нельзя удалить собственный аккаунт")
                    deleted_name = delete_user(target["id"])
                    message = f"Пользователь {deleted_name} удалён"
                else:
                    raise ValueError("Неизвестное действие")
        except ValueError as operation_error:
            return redirect(url_for("admin_users", error=str(operation_error)))
        return redirect(url_for("admin_users", message=message))

    return render_template(
        "settings.html",
        mode="users",
        title="Пользователи",
        subtitle="Аккаунты, роли и доступ к Монитору",
        current_user=administrator,
        users=list_users(),
        csrf_token=csrf_token(),
        message=str(request.args.get("message", "")).strip(),
        error=str(request.args.get("error", "")).strip(),
    )


def _registered_admin_sources():
    """Возвращает источники интерфейса без старых совместимых псевдонимов."""
    groups = (
        ("Госструктуры", GOVERNMENT_SOURCES - {"Сахалинская обл."}),
        ("Информагентства", AGENCY_SOURCES),
        ("Газеты", NEWSPAPER_SOURCES),
    )
    return [
        (source, group_label)
        for group_label, names in groups
        for source in sorted(names, key=str.casefold)
    ]


def _format_file_size(value):
    """Показывает размер файла без технических байтовых чисел."""
    try:
        size = max(0, int(value))
    except (TypeError, ValueError):
        size = 0
    units = ("Б", "КБ", "МБ", "ГБ")
    amount = float(size)
    for unit in units:
        if amount < 1024 or unit == units[-1]:
            return f"{amount:.0f} {unit}" if unit == "Б" else f"{amount:.1f} {unit}"
        amount /= 1024
    return f"{size} Б"


def _format_duration(value):
    """Показывает длительность инцидента компактно и без секундного шума."""
    try:
        seconds = max(0, int(value))
    except (TypeError, ValueError):
        seconds = 0
    days, remainder = divmod(seconds, 86400)
    hours, remainder = divmod(remainder, 3600)
    minutes = remainder // 60
    if days:
        return f"{days} д. {hours} ч."
    if hours:
        return f"{hours} ч. {minutes} мин."
    return f"{max(1, minutes)} мин."


@app.route("/admin/sources", methods=["GET", "POST"])
def admin_sources():
    """Показывает состояние парсеров и ставит ручные проверки в очередь."""
    administrator = current_user()
    if not administrator or administrator.get("role") != "admin":
        abort(403)

    registered = _registered_admin_sources()
    registered_names = {source for source, _group in registered}
    settings = load_source_settings()

    if request.method == "POST":
        if not csrf_is_valid():
            abort(400)
        source = " ".join(str(request.form.get("source", "")).split())
        if source not in registered_names:
            return redirect(url_for("admin_sources", error="Источник не найден"))
        action = str(request.form.get("action", "")).strip()
        try:
            if action == "toggle":
                enabled = settings.get(source, {}).get("enabled", True)
                changed = set_source_enabled(source, not enabled)
                state = "возвращён в расписание" if changed["enabled"] else "поставлен на паузу"
                message = f"{source}: {state}"
            elif action == "refresh":
                if not settings.get(source, {}).get("enabled", True):
                    raise ValueError("Сначала включите источник")
                job = enqueue_parser_job(source, administrator.get("id"))
                state = "уже выполняется" if job["status"] == "running" else "поставлен в очередь"
                message = f"{source}: {state}"
            else:
                raise ValueError("Неизвестное действие")
        except ValueError as operation_error:
            return redirect(url_for("admin_sources", error=str(operation_error)))
        return redirect(url_for("admin_sources", message=message))

    status_document = load_json("parser_status.json", {})
    status_by_source = {
        item.get("source"): item
        for item in status_document.get("sources", [])
        if item.get("source")
    }
    database_by_source = source_news_statistics()
    jobs = list_parser_jobs(30)
    latest_job_by_source = {}
    for job in jobs:
        latest_job_by_source.setdefault(job["source"], job)

    status_labels = {
        "ok": "Работает",
        "empty": "Пустая выдача",
        "error": "Ошибка",
        "disabled": "На паузе",
        "unknown": "Нет данных",
        "pending": "В очереди",
        "running": "Обновляется",
    }
    job_labels = {
        "pending": "В очереди",
        "running": "Выполняется",
        "success": "Выполнено",
        "error": "Ошибка",
    }
    prepared_sources = []
    for source, group_label in registered:
        parser_status = status_by_source.get(source, {})
        database_status = database_by_source.get(source, {})
        enabled = settings.get(source, {}).get("enabled", True)
        job = latest_job_by_source.get(source, {})
        job_active = job.get("status") in {"pending", "running"}
        raw_status = parser_status.get("status", "unknown") if enabled else "disabled"
        display_status = job.get("status") if enabled and job_active else raw_status
        error = (
            job.get("error", "")
            if job.get("status") == "error"
            else parser_status.get("error", "")
        )
        prepared_sources.append({
            "source": source,
            "group_label": group_label,
            "enabled": enabled,
            "status_class": display_status,
            "status_label": status_labels.get(display_status, display_status),
            "job_active": job_active,
            "job_label": job_labels.get(job.get("status"), "") if job else "",
            "checked_at": parser_status.get("checked_at", ""),
            "last_success": parser_status.get("last_success", ""),
            "last_received": database_status.get("last_received", ""),
            "newest_publication": database_status.get("newest_publication", ""),
            "total_news": database_status.get("news_count", 0),
            "news_count": parser_status.get("news_count", 0),
            "duration": parser_status.get("duration_seconds", 0),
            "failure_streak": parser_status.get("failure_streak", 0),
            "error": error,
        })

    enabled_sources = [item for item in prepared_sources if item["enabled"]]
    summary = {
        "total": len(prepared_sources),
        "ok": sum(item["status_class"] == "ok" for item in enabled_sources),
        "problem": sum(
            item["status_class"] in {"empty", "error", "unknown"}
            for item in enabled_sources
        ),
        "disabled": len(prepared_sources) - len(enabled_sources),
    }
    prepared_jobs = []
    for job in jobs:
        prepared = dict(job)
        prepared["status_label"] = job_labels.get(job["status"], job["status"])
        prepared_jobs.append(prepared)

    alerts = source_alerts(prepared_sources)

    return render_template(
        "admin_sources.html",
        sources=prepared_sources,
        jobs=prepared_jobs,
        summary=summary,
        alerts=alerts,
        alert_counts=alert_summary(alerts),
        kyodo_vpn=kyodo_proxy_status(),
        auto_refresh=any(job["status"] in {"pending", "running"} for job in jobs),
        current_user=administrator,
        csrf_token=csrf_token(),
        message=str(request.args.get("message", "")).strip(),
        error=str(request.args.get("error", "")).strip(),
    )


@app.route("/admin/system", methods=["GET", "POST"])
def admin_system():
    """Показывает состояние SQLite, копии базы и хвост журнала ошибок."""
    administrator = current_user()
    if not administrator or administrator.get("role") != "admin":
        abort(403)

    if request.method == "POST":
        if not csrf_is_valid():
            abort(400)
        action = str(request.form.get("action", "")).strip()
        if action != "backup":
            return redirect(url_for("admin_system", error="Неизвестное действие"))
        try:
            created = create_manual_backup(retention=10)
        except Exception as backup_error:
            return redirect(url_for(
                "admin_system",
                error=(
                    "Не удалось создать резервную копию: "
                    f"{type(backup_error).__name__}: {backup_error}"
                ),
            ))
        return redirect(url_for(
            "admin_system",
            message=f"Создана резервная копия {created['name']}",
        ))

    database = database_stats()
    prepared_database = dict(database)
    prepared_database["size"] = _format_file_size(database.get("size_bytes", 0))
    backups = list_database_backups()
    prepared_backups = []
    for item in backups:
        prepared = dict(item)
        prepared["size"] = _format_file_size(item.get("size_bytes", 0))
        prepared_backups.append(prepared)
    log = error_log_stats()
    log["size"] = _format_file_size(log.get("size_bytes", 0))
    alerts = system_alerts(database, backups)

    return render_template(
        "admin_system.html",
        database=prepared_database,
        backups=prepared_backups,
        errors=list(reversed(read_recent_errors(limit=120))),
        log=log,
        alerts=alerts,
        alert_counts=alert_summary(alerts),
        version=PROJECT_VERSION,
        current_user=administrator,
        csrf_token=csrf_token(),
        message=str(request.args.get("message", "")).strip(),
        error=str(request.args.get("error", "")).strip(),
    )


@app.route("/admin/incidents")
def admin_incidents():
    """Показывает администраторам историю сбоев и восстановлений."""
    administrator = current_user()
    if not administrator or administrator.get("role") != "admin":
        abort(403)
    state = str(request.args.get("state", "all")).strip().casefold()
    if state not in {"all", "active", "resolved"}:
        state = "all"
    prepared = []
    for incident in list_source_incidents(state=state, limit=250):
        item = dict(incident)
        item["duration"] = _format_duration(item.get("duration_seconds", 0))
        item["status_label"] = (
            "Критично"
            if item["is_active"] and item["level"] == "critical"
            else ("Внимание" if item["is_active"] else "Восстановлен")
        )
        prepared.append(item)
    return render_template(
        "admin_incidents.html",
        incidents=prepared,
        summary=source_incident_statistics(),
        state=state,
        current_user=administrator,
    )


@app.route("/admin/reliability")
def admin_reliability():
    """Показывает администраторам доступность источников за выбранный период."""
    administrator = current_user()
    if not administrator or administrator.get("role") != "admin":
        abort(403)
    try:
        days = int(request.args.get("days", 7))
    except (TypeError, ValueError):
        days = 7
    if days not in {7, 30}:
        days = 7

    registered = _registered_admin_sources()
    group_by_source = dict(registered)
    settings = load_source_settings()
    statistics = source_reliability_statistics(
        days=days,
        sources=[source for source, _group in registered],
    )
    prepared = []
    for raw_item in statistics:
        item = dict(raw_item)
        uptime = item["uptime_percent"]
        item["group_label"] = group_by_source.get(item["source"], "Другие")
        item["enabled"] = settings.get(item["source"], {}).get("enabled", True)
        item["uptime_display"] = f"{uptime:.3f}".rstrip("0").rstrip(".")
        item["uptime_class"] = (
            "bad" if uptime < 95 else ("warn" if uptime < 99.5 else "good")
        )
        item["downtime"] = (
            _format_duration(item["downtime_seconds"])
            if item["downtime_seconds"]
            else "—"
        )
        item["average_duration"] = (
            _format_duration(item["average_incident_seconds"])
            if item["average_incident_seconds"]
            else "—"
        )
        prepared.append(item)

    total = len(prepared)
    average_uptime = (
        sum(item["uptime_percent"] for item in prepared) / total
        if total
        else 100
    )
    total_downtime = sum(item["downtime_seconds"] for item in prepared)
    return render_template(
        "admin_reliability.html",
        sources=prepared,
        days=days,
        summary={
            "average_uptime": f"{average_uptime:.3f}".rstrip("0").rstrip("."),
            "incidents": sum(item["incident_count"] for item in prepared),
            "downtime": (
                _format_duration(total_downtime) if total_downtime else "0 мин."
            ),
            "affected": sum(item["incident_count"] > 0 for item in prepared),
            "total": total,
        },
        current_user=administrator,
    )


def _matches_collection_search(item, query, fields):
    if not query:
        return True
    haystack = "\n".join(str(item.get(field, "") or "") for field in fields)
    return query.casefold() in haystack.casefold()


def _prepare_collection_note(note):
    """Оставляет в ленте первый абзац или до 650 знаков."""
    prepared = dict(note)
    body = str(prepared.get("body", "") or "").strip()
    paragraph_break = re.search(r"\n\s*\n", body)
    preview_limit = min(paragraph_break.start() if paragraph_break else len(body), 650)
    split_at = preview_limit
    if preview_limit == 650 and len(body) > preview_limit:
        word_break = body.rfind(" ", 0, preview_limit)
        if word_break >= 400:
            split_at = word_break
    prepared["body_preview"] = body[:split_at].rstrip()
    prepared["body_remainder"] = body[split_at:] if split_at < len(body) else ""
    return prepared


COLLECTION_SORTS = {
    "newest": "Сначала новые",
    "oldest": "Сначала старые",
    "title": "По заголовку",
    "source": "По источнику",
}


def _collection_materials(bookmarks, notes, sort_mode="newest"):
    """Объединяет статьи из ленты и добавленные вручную материалы для сортировки."""
    materials = []
    for note in notes:
        item = dict(note)
        item["kind"] = "note"
        item["date"] = item.get("publication_date", "")
        materials.append(item)
    for bookmark in bookmarks:
        item = dict(bookmark)
        item["kind"] = "bookmark"
        item["publication_date"] = item.get("date", "")
        materials.append(item)

    sort_mode = sort_mode if sort_mode in COLLECTION_SORTS else "newest"
    if sort_mode in {"title", "source"}:
        field = sort_mode
        return sorted(
            materials,
            key=lambda item: (
                str(item.get(field, "") or "").casefold(),
                str(item.get("title", "") or "").casefold(),
            ),
        )

    def date_key(item):
        return (
            bool(item.get("publication_date")),
            str(item.get("publication_date", "") or ""),
            str(item.get("updated_at", "") or item.get("created_at", "") or ""),
            int(item.get("id", 0)),
        )

    return sorted(materials, key=date_key, reverse=sort_mode == "newest")


def _collection_export_name(value):
    safe = re.sub(r"[^0-9A-Za-zА-Яа-яЁё_-]+", "-", str(value or "")).strip("-")
    return (safe or "podborka")[:80] + ".docx"


@app.get("/collections/export.docx")
def export_collection_docx():
    """Скачивает доступную подборку как Word-документ со ссылками и реквизитами."""
    from docx import Document
    from docx.shared import Pt

    user = current_user()
    folder_id = request.args.get("folder")
    collection = load_collection(user["id"], folder_id)
    if collection is None:
        abort(404)
    materials = _collection_materials(
        list_collection_bookmarks(user["id"], folder_id),
        list_collection_notes(user["id"], folder_id),
        request.args.get("sort", "newest"),
    )

    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(11)
    document.add_heading(collection["name"], level=0)
    if collection.get("description"):
        document.add_paragraph(collection["description"])
    document.add_paragraph(f"Материалов: {len(materials)}")
    for number, item in enumerate(materials, start=1):
        document.add_heading(f"{number}. {item.get('title') or 'Без заголовка'}", level=1)
        document.add_paragraph(f"Источник: {item.get('source') or 'не указан'}")
        document.add_paragraph(
            f"Дата публикации: {item.get('publication_date') or 'не указана'}"
        )
        document.add_paragraph(f"Ссылка: {item.get('url') or 'не указана'}")

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return send_file(
        output,
        as_attachment=True,
        download_name=_collection_export_name(collection["name"]),
        mimetype=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )


MONTH_NAMES_RU = (
    "", "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
    "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь",
)
ACCESS_LABELS = {
    "private": "Только я", "selected": "Выбранные пользователи",
    "all": "Все пользователи",
}


def _notes_admin():
    user = current_user()
    if not user or user.get("role") != "admin":
        abort(403)
    return user


def _notes_redirect(view, **values):
    return redirect(url_for("notes_page", view=view, **values))


@app.route("/notes", methods=["GET", "POST"])
def notes_page():
    """Экспериментальные заметки, доступные только администратору."""
    user = _notes_admin()
    user_id = user["id"]
    view = str(request.values.get("view", "calendar")).strip().casefold()
    if view not in {"calendar", "records", "dictionary"}:
        view = "calendar"

    if request.method == "POST":
        if not csrf_is_valid():
            abort(400)
        action = str(request.form.get("action", "")).strip()
        try:
            if action == "save_event":
                event_id = save_calendar_event(
                    user_id, request.form.get("title"),
                    request.form.get("event_date"), request.form.get("event_time"),
                    request.form.get("place"), request.form.get("description"),
                    request.form.get("visibility"),
                    request.form.getlist("shared_user_ids"),
                    request.form.get("event_id"),
                )
                event_date = request.form.get("event_date")
                return _notes_redirect(
                    "calendar", selected=event_date, event=event_id,
                    message="Мероприятие сохранено",
                )
            if action == "delete_event":
                delete_calendar_event(user_id, request.form.get("event_id"))
                return _notes_redirect(
                    "calendar", selected=request.form.get("return_date"),
                    message="Мероприятие удалено",
                )
            if action == "save_note":
                if "delete_note" in request.form.getlist("action"):
                    delete_personal_note(user_id, request.form.get("note_id"))
                    return _notes_redirect("records", message="Запись удалена")
                note_id = save_personal_note(
                    user_id, request.form.get("folder"), request.form.get("title"),
                    request.form.get("body"), request.form.get("visibility"),
                    request.form.getlist("shared_user_ids"),
                    request.form.get("note_id"),
                )
                return _notes_redirect(
                    "records", note=note_id, message="Запись сохранена"
                )
            if action == "create_deck":
                deck_id = create_dictionary_deck(user_id, request.form.get("name"))
                return _notes_redirect(
                    "dictionary", deck=deck_id, message="Словарь создан"
                )
            if action == "add_card":
                deck_id = request.form.get("deck_id")
                save_dictionary_card(
                    user_id, deck_id, request.form.get("term"),
                    request.form.get("reading"), request.form.get("translation"),
                )
                return _notes_redirect(
                    "dictionary", deck=deck_id, message="Карточка добавлена"
                )
            if action == "review_card":
                deck_id = request.form.get("deck_id")
                result = review_dictionary_card(
                    user_id, request.form.get("card_id"), request.form.get("rating")
                )
                return _notes_redirect(
                    "dictionary", deck=deck_id,
                    message=f"Следующий повтор: {result['next_review']}",
                )
            raise ValueError("Неизвестное действие")
        except ValueError as operation_error:
            return _notes_redirect(view, error=str(operation_error))

    available_users = [
        account for account in list_users()
        if account["is_active"] and account["id"] != user_id
    ]
    context = {
        "current_user": user,
        "csrf_token": csrf_token(),
        "view": view,
        "available_users": available_users,
        "message": str(request.args.get("message", "")).strip(),
        "error": str(request.args.get("error", "")).strip(),
        "access_labels": ACCESS_LABELS,
    }

    if view == "calendar":
        today = date.today()
        try:
            year = int(request.args.get("year", today.year))
            month = int(request.args.get("month", today.month))
            if not 2000 <= year <= 2100 or not 1 <= month <= 12:
                raise ValueError
        except (TypeError, ValueError):
            year, month = today.year, today.month
        first_day = date(year, month, 1)
        days_in_month = monthrange(year, month)[1]
        last_day = date(year, month, days_in_month)
        grid_start = first_day - timedelta(days=first_day.weekday())
        grid_end = grid_start + timedelta(days=41)
        events = list_calendar_events(
            user_id, grid_start.isoformat(), grid_end.isoformat()
        )
        events_by_date = {}
        for event in events:
            events_by_date.setdefault(event["event_date"], []).append(event)
        selected_date = str(request.args.get("selected", "")).strip()
        try:
            parsed_selected = datetime.strptime(selected_date, "%Y-%m-%d").date()
        except ValueError:
            parsed_selected = today if first_day <= today <= last_day else first_day
            selected_date = parsed_selected.isoformat()
        calendar_days = []
        for offset in range(42):
            item_date = grid_start + timedelta(days=offset)
            calendar_days.append({
                "number": item_date.day,
                "iso": item_date.isoformat(),
                "in_month": item_date.month == month,
                "events": events_by_date.get(item_date.isoformat(), []),
                "url": url_for(
                    "notes_page", view="calendar", year=item_date.year,
                    month=item_date.month, selected=item_date.isoformat(),
                ),
            })
        selected_events = list(events_by_date.get(selected_date, []))
        for event in selected_events:
            event["edit_url"] = url_for(
                "notes_page", view="calendar", year=year, month=month,
                selected=selected_date, event=event["id"],
            )
        selected_event = next(
            (
                event for event in events
                if str(event["id"]) == str(request.args.get("event", ""))
            ),
            None,
        )
        event_form = selected_event or {
            "id": "", "title": "", "event_date": selected_date,
            "event_time": "", "place": "", "description": "",
            "visibility": "private", "shared_users": [],
        }
        previous_anchor = first_day - timedelta(days=1)
        next_anchor = last_day + timedelta(days=1)
        context.update(
            month_label=f"{MONTH_NAMES_RU[month]} {year}",
            calendar_days=calendar_days,
            selected_date=selected_date,
            selected_date_label=parsed_selected.strftime("%d.%m.%Y"),
            selected_events=selected_events,
            event_form=event_form,
            event_shared_ids={item["id"] for item in event_form["shared_users"]},
            previous_month_url=url_for(
                "notes_page", view="calendar", year=previous_anchor.year,
                month=previous_anchor.month,
            ),
            next_month_url=url_for(
                "notes_page", view="calendar", year=next_anchor.year,
                month=next_anchor.month,
            ),
            today_url=url_for(
                "notes_page", view="calendar", year=today.year,
                month=today.month, selected=today.isoformat(),
            ),
            new_event_url=url_for(
                "notes_page", view="calendar", year=year, month=month,
                selected=selected_date,
            ),
        )
    elif view == "records":
        notes = list_personal_notes(user_id)
        selected_folder = str(request.args.get("folder", "")).strip()
        folder_counts = {}
        for note in notes:
            folder_counts[note["folder"]] = folder_counts.get(note["folder"], 0) + 1
            note["url"] = url_for(
                "notes_page", view="records", folder=selected_folder or None,
                note=note["id"],
            )
        filtered_notes = [
            note for note in notes
            if not selected_folder or note["folder"] == selected_folder
        ]
        selected_note = next(
            (
                note for note in notes
                if str(note["id"]) == str(request.args.get("note", ""))
            ),
            None,
        )
        note_form = selected_note or {
            "id": "", "folder": selected_folder or "Без папки", "title": "",
            "body": "", "visibility": "private", "shared_users": [],
        }
        context.update(
            notes=notes,
            filtered_notes=filtered_notes,
            selected_folder=selected_folder,
            selected_note=selected_note,
            note_form=note_form,
            note_shared_ids={item["id"] for item in note_form["shared_users"]},
            note_folders=[
                {"name": name, "count": count,
                 "url": url_for("notes_page", view="records", folder=name)}
                for name, count in sorted(folder_counts.items())
            ],
        )
    else:
        decks = list_dictionary_decks(user_id)
        requested_deck = str(request.args.get("deck", "")).strip()
        selected_deck = next(
            (deck for deck in decks if str(deck["id"]) == requested_deck),
            decks[0] if decks else None,
        )
        for deck in decks:
            deck["url"] = url_for("notes_page", view="dictionary", deck=deck["id"])
        cards = (
            list_dictionary_cards(user_id, selected_deck["id"])
            if selected_deck else []
        )
        due_cards = (
            list_dictionary_cards(user_id, selected_deck["id"], due_only=True)
            if selected_deck else []
        )
        context.update(
            decks=decks,
            deck_total=len(decks),
            due_total=sum(deck["due_count"] for deck in decks),
            selected_deck=selected_deck,
            cards=cards,
            quiz_card=due_cards[0] if due_cards else None,
            ratings=(("again", "Снова"), ("hard", "Трудно"),
                     ("good", "Хорошо"), ("easy", "Легко")),
        )
    return render_template("notes.html", **context)


@app.route("/bookmarks", methods=["GET", "POST"])
@app.route("/collections", methods=["GET", "POST"])
def bookmarks_page():
    """Показывает рабочие подборки, общий доступ, ссылки и заметки."""
    user = current_user()
    user_id = user["id"]
    favorite_folder = ensure_favorites_folder(user_id)
    selected_folder = str(
        request.args.get("folder", favorite_folder["id"])
    ).strip() or str(favorite_folder["id"])
    sort_mode = str(request.args.get("sort", "newest")).strip().casefold()
    if sort_mode not in COLLECTION_SORTS:
        sort_mode = "newest"
    folders = list_bookmark_folders(user_id)
    folder_by_id = {str(folder["id"]): folder for folder in folders}
    shared_folders = list_shared_collections(user_id)
    shared_by_id = {str(folder["id"]): folder for folder in shared_folders}
    if selected_folder not in {"all", "unfiled", *folder_by_id, *shared_by_id}:
        abort(404)

    if request.method == "POST":
        if not csrf_is_valid():
            abort(400)
        action = str(request.form.get("action", "")).strip()
        try:
            if action == "create_folder":
                created = create_bookmark_folder(user_id, request.form.get("name"))
                message = f"Подборка «{created['name']}» создана"
                selected_folder = str(created["id"])
            elif action == "update_collection":
                current = load_collection(user_id, request.form.get("folder_id"))
                if current is None or not current["can_edit"]:
                    raise ValueError("Подборка не найдена")
                updated = update_collection(
                    user_id,
                    request.form.get("folder_id"),
                    request.form.get("name"),
                    current["description"],
                    current["visibility"],
                    [item["id"] for item in current["shared_users"]],
                )
                message = f"Название изменено на «{updated['name']}»"
                selected_folder = str(updated["id"])
            elif action == "share_collection":
                current = load_collection(user_id, request.form.get("folder_id"))
                if current is None or not current["can_edit"]:
                    raise ValueError("Подборка не найдена")
                update_collection(
                    user_id, current["id"], current["name"],
                    request.form.get("description"),
                    request.form.get("visibility"),
                    request.form.getlist("shared_user_ids"),
                )
                message = "Доступ к подборке обновлён"
                selected_folder = str(current["id"])
            elif action == "delete_folder":
                delete_bookmark_folder(user_id, request.form.get("folder_id"))
                message = "Подборка удалена, её новости перенесены в «Без подборки»"
                selected_folder = str(favorite_folder["id"])
            elif action == "add_external":
                save_external_bookmark(
                    user_id, request.form.get("folder_id"), request.form.get("url"),
                    request.form.get("title"), request.form.get("note"),
                )
                message = "Внешняя ссылка добавлена"
                selected_folder = str(request.form.get("folder_id"))
            elif action == "add_note":
                save_collection_note(
                    user_id, request.form.get("folder_id"),
                    request.form.get("title"), request.form.get("body"),
                    request.form.get("url"), request.form.get("source"),
                    request.form.get("publication_date"),
                    request.form.get("comment"),
                )
                message = "Статья добавлена"
                selected_folder = str(request.form.get("folder_id"))
            elif action == "update_note":
                update_collection_note(
                    user_id, request.form.get("folder_id"),
                    request.form.get("note_id"), request.form.get("title"),
                    request.form.get("body"), request.form.get("url"),
                    request.form.get("source"), request.form.get("publication_date"),
                    request.form.get("comment"),
                )
                message = "Изменения сохранены"
                selected_folder = str(request.form.get("folder_id"))
            elif action == "delete_note":
                delete_collection_note(
                    user_id, request.form.get("folder_id"), request.form.get("note_id"),
                )
                message = "Статья удалена"
                selected_folder = str(request.form.get("folder_id"))
            elif action == "update_bookmark":
                update_bookmark(
                    user_id,
                    request.form.get("bookmark_url"),
                    request.form.get("folder_id"),
                    request.form.get("note"),
                )
                message = "Папка и заметка сохранены"
            elif action == "remove_bookmark":
                remove_bookmark(user_id, request.form.get("bookmark_url"))
                message = "Новость удалена из закладок"
            else:
                raise ValueError("Неизвестное действие")
        except ValueError as operation_error:
            return redirect(
                url_for(
                    "bookmarks_page",
                    folder=selected_folder,
                    sort=sort_mode,
                    error=str(operation_error),
                )
            )
        return redirect(url_for(
            "bookmarks_page", folder=selected_folder, sort=sort_mode, message=message
        ))

    search_query = str(request.args.get("q", "")).strip()[:200]
    all_bookmarks = list_bookmarks(user_id)
    selected_folder_data = None
    notes = []
    read_note_ids = set()
    if selected_folder in folder_by_id or selected_folder in shared_by_id:
        selected_folder_data = load_collection(user_id, selected_folder)
        bookmarks = list_collection_bookmarks(user_id, selected_folder)
        notes = list_collection_notes(user_id, selected_folder)
        read_note_ids = list_collection_note_read_ids(user_id, selected_folder)
    else:
        bookmarks = list_bookmarks(user_id, selected_folder)
    bookmarks = [
        item for item in bookmarks
        if _matches_collection_search(
            item,
            search_query,
            ("title", "source", "note", "date", "url", "folder_name"),
        )
    ]
    prepared_notes = []
    for item in notes:
        if not _matches_collection_search(
            item,
            search_query,
            ("title", "source", "body", "comment", "publication_date", "url"),
        ):
            continue
        prepared = _prepare_collection_note(item)
        prepared["is_read"] = int(item["id"]) in read_note_ids
        prepared_notes.append(prepared)
    notes = prepared_notes
    materials = _collection_materials(bookmarks, notes, sort_mode)
    notes = [item for item in materials if item["kind"] == "note"]
    bookmarks = [item for item in materials if item["kind"] == "bookmark"]
    if selected_folder_data:
        selected_title = selected_folder_data["name"]
    elif selected_folder == "unfiled":
        selected_title = "Без папки"
    else:
        selected_title = "Все сохранённые"
    active_users = [
        account for account in list_users()
        if account["is_active"] and account["id"] != user_id
    ]
    return render_template(
        "bookmarks.html",
        current_user=user,
        csrf_token=csrf_token(),
        folders=folders,
        shared_folders=shared_folders,
        bookmarks=bookmarks,
        notes=notes,
        materials=materials,
        selected_folder=selected_folder,
        selected_folder_data=selected_folder_data,
        selected_title=selected_title,
        search_query=search_query,
        sort_mode=sort_mode,
        sort_options=COLLECTION_SORTS,
        total_count=len(all_bookmarks),
        unfiled_count=sum(item["folder_id"] is None for item in all_bookmarks),
        available_users=active_users,
        selected_shared_ids={
            account["id"] for account in (
                selected_folder_data["shared_users"] if selected_folder_data else []
            )
        },
        message=str(request.args.get("message", "")).strip(),
        error=str(request.args.get("error", "")).strip(),
    )


def can_view_admin_diagnostics():
    """Служебная диагностика видна только администратору."""
    user = current_user()
    return bool(user and user.get("role") == "admin")


def apply_source_order(sources, preferred_names):
    """Ставит сохранённые пользователем источники первыми, новые — в конец."""
    remaining = {
        str(name).casefold(): (name, count)
        for name, count in sources
    }
    ordered = []
    for preferred in preferred_names or []:
        item = remaining.pop(str(preferred).casefold(), None)
        if item is not None:
            ordered.append(item)
    ordered.extend(
        item
        for item in sources
        if str(item[0]).casefold() in remaining
    )
    return ordered


def render_news_page(
    mode="all",
    source_filters=None,
    source_group=GOVERNMENT_GROUP,
):
    user = current_user()
    if user["id"]:
        user_saved_urls = bookmarked_urls(user["id"])
        user_bookmark_count = count_bookmarks(user["id"])
    else:
        user_saved_urls = []
        user_bookmark_count = 0
    status = load_json("parser_status.json", {})
    total, found_count = news_group_counts(source_group)
    counts = news_source_counts(source_group)
    sources = sorted(counts.items(), key=lambda item: (-item[1], item[0]))
    requested_sources = (
        list(source_filters)
        if source_filters is not None
        else request.args.getlist("source")
    )
    available_names = set(counts)
    source_filters = []
    for source in requested_sources:
        source = str(source or "").strip()
        if source in available_names and source not in source_filters:
            source_filters.append(source)
    if user["id"]:
        sources = apply_source_order(
            sources,
            load_source_order(user["id"], source_group),
        )
    if source_group == AGENCIES_GROUP:
        yahoo_sources = [
            (
                name,
                count,
                name.split("·", 1)[-1].strip() or name,
            )
            for name, count in sources
            if is_yahoo_source(name)
        ]
        sidebar_sources = [
            (name, count)
            for name, count in sources
            if not is_yahoo_source(name)
        ]
    else:
        yahoo_sources = []
        sidebar_sources = sources
    yahoo_active = any(is_yahoo_source(source) for source in source_filters)
    yahoo_expanded = yahoo_active

    status_sources = [
        item
        for item in status.get("sources", [])
        if get_source_group(item.get("source", "")) == source_group
    ]
    total_sources = len(status_sources) or len(sources)
    ok_sources = sum(
        item.get("status") == "ok"
        for item in status_sources
    )
    if not status_sources:
        ok_sources = total_sources
    problem_sources = sorted(
        (
            item
            for item in status_sources
            if item.get("status") in {"empty", "error"}
        ),
        key=lambda item: (
            item.get("status") != "error",
            item.get("source", ""),
        ),
    )

    if source_group == AGENCIES_GROUP:
        group_title = "Новости информагентств"
        group_eyebrow = (
            "РИА Новости · ТАСС · Интерфакс · Yahoo! JAPAN · "
            "Yonhap · Киодо"
        )
        group_home = "/agencies"
        group_found = "/agencies/found"
    elif source_group == NEWSPAPERS_GROUP:
        group_title = "Свежие номера газет"
        group_eyebrow = "Коммерсантъ · Известия · РГ · Ведомости · Красная звезда · КП"
        group_home = "/newspapers"
        group_found = "/newspapers/found"
    else:
        group_title = "Новости госструктур"
        group_eyebrow = "Агрегатор официальных источников"
        group_home = "/"
        group_found = "/found"

    filter_home = group_found if mode == "found" else group_home
    requested_keyword = " ".join(request.args.get("keyword", "").split())
    keyword_filter = (
        requested_keyword[:80]
        if mode == "found"
        else ""
    )
    shared_query_parameters = [
        (key, value)
        for key, values in request.args.lists()
        if key not in {"page", "source", "keyword"}
        for value in values
    ]
    persistent_query_parameters = list(shared_query_parameters)
    if keyword_filter:
        persistent_query_parameters.append(("keyword", keyword_filter))
    clear_query = urlencode(persistent_query_parameters, doseq=True)
    clear_sources_url = filter_home + (f"?{clear_query}" if clear_query else "")
    source_query = urlencode([("source", source) for source in source_filters])
    group_home_url = group_home + (f"?{source_query}" if source_query else "")
    found_query_parameters = [
        *(('source', source) for source in source_filters),
        *((('keyword', keyword_filter),) if keyword_filter else ()),
    ]
    found_query = urlencode(found_query_parameters, doseq=True)
    group_found_url = group_found + (f"?{found_query}" if found_query else "")
    if mode == "found":
        feed_title = (
            f"Совпадения: {keyword_filter}"
            if keyword_filter
            else "Совпадения"
        )
    elif len(source_filters) == 1:
        feed_title = f"Источник: {source_filters[0]}"
    elif source_filters:
        feed_title = f"Выбрано источников: {len(source_filters)}"
    else:
        feed_title = "Последние публикации"

    search_query = request.args.get("q", "").strip()
    try:
        page = max(1, int(request.args.get("page", "1")))
    except (TypeError, ValueError):
        page = 1
    page_offset = (page - 1) * NEWS_PER_PAGE
    page_news, page_total = list_news_page(
        source_group,
        found_only=mode == "found",
        sources=source_filters,
        search_query=search_query,
        keyword=keyword_filter,
        limit=NEWS_PER_PAGE,
        offset=page_offset,
    )
    page_count = max(1, (page_total + NEWS_PER_PAGE - 1) // NEWS_PER_PAGE)
    if page > page_count:
        page = page_count
        page_offset = (page - 1) * NEWS_PER_PAGE
        page_news, page_total = list_news_page(
            source_group,
            found_only=mode == "found",
            sources=source_filters,
            search_query=search_query,
            keyword=keyword_filter,
            limit=NEWS_PER_PAGE,
            offset=page_offset,
        )
    page_start = page_offset + 1 if page_news else 0
    page_end = page_offset + len(page_news)
    page_label = (
        f"{page_start}–{page_end} из {page_total}"
        if page_total
        else "0 материалов"
    )

    def page_url(number):
        parameters = list(persistent_query_parameters)
        parameters.extend(("source", source) for source in source_filters)
        if number > 1:
            parameters.append(("page", number))
        query_string = urlencode(parameters, doseq=True)
        return request.path + (f"?{query_string}" if query_string else "")

    visible_pages = sorted(
        {1, page_count}
        | set(range(max(1, page - 2), min(page_count, page + 2) + 1))
    )
    page_links = []
    previous_number = None
    for number in visible_pages:
        if previous_number is not None and number - previous_number > 1:
            page_links.append({"gap": True})
        page_links.append(
            {
                "gap": False,
                "number": number,
                "url": page_url(number),
                "current": number == page,
            }
        )
        previous_number = number

    keyword_urls = {}
    for item in page_news:
        for keyword in item.get("keywords", []) or []:
            parameters = list(shared_query_parameters)
            parameters.extend(("source", source) for source in source_filters)
            parameters.append(("keyword", keyword))
            query_string = urlencode(parameters, doseq=True)
            keyword_urls[keyword] = group_found + f"?{query_string}"

    unread_summary = news_unread_summary(
        user["id"],
        source_group,
        [item.get("url", "") for item in page_news],
    )
    unread_counts = unread_summary["by_source"]

    return render_template(
        "news.html",
        news=page_news,
        total=total,
        found_count=found_count,
        sources=sources,
        sidebar_sources=sidebar_sources,
        yahoo_sources=yahoo_sources,
        yahoo_active=yahoo_active,
        yahoo_expanded=yahoo_expanded,
        keyword_filter=keyword_filter,
        keyword_urls=keyword_urls,
        unread_urls=unread_summary["visible_urls"],
        unread_counts=unread_counts,
        unread_total=unread_summary["total"],
        yahoo_unread_count=sum(
            count
            for source, count in unread_counts.items()
            if is_yahoo_source(source)
        ),
        source_filters=source_filters,
        feed_title=feed_title,
        mode=mode,
        source_group=source_group,
        group_title=group_title,
        group_eyebrow=group_eyebrow,
        group_home=group_home,
        group_found=group_found,
        group_home_url=group_home_url,
        group_found_url=group_found_url,
        filter_home=filter_home,
        clear_sources_url=clear_sources_url,
        health_total=total_sources,
        health_ok=ok_sources,
        health_empty=sum(
            item.get("status") == "empty"
            for item in status_sources
        ),
        health_errors=sum(
            item.get("status") == "error"
            for item in status_sources
        ),
        problem_sources=problem_sources,
        status_time=status.get("generated_at", ""),
        current_path=request.path,
        current_url=request.full_path.rstrip("?"),
        search_query=search_query,
        page=page,
        page_count=page_count,
        page_links=page_links,
        previous_url=page_url(page - 1) if page > 1 else "",
        next_url=page_url(page + 1) if page < page_count else "",
        page_label=page_label,
        show_admin_diagnostics=can_view_admin_diagnostics(),
        current_user=user,
        csrf_token=csrf_token(),
        saved_urls=user_saved_urls,
        bookmark_count=user_bookmark_count,
    )


@app.template_filter("urlencode")
def urlencode_filter(value):
    return quote(str(value), safe="")


@app.route("/")
def index():
    return render_news_page(
        source_group=GOVERNMENT_GROUP,
    )


@app.route("/found")
def found_page():
    return render_news_page(
        mode="found",
        source_group=GOVERNMENT_GROUP,
    )


@app.route("/filter/<path:source>")
def filter_source(source):
    return render_news_page(
        source_filters=[source],
        source_group=get_source_group(source),
    )


@app.route("/agencies")
def agencies_page():
    return render_news_page(
        source_group=AGENCIES_GROUP,
    )


@app.route("/agencies/found")
def agencies_found_page():
    return render_news_page(
        mode="found",
        source_group=AGENCIES_GROUP,
    )


@app.route("/agencies/filter/<path:source>")
def agencies_filter_source(source):
    return render_news_page(
        source_filters=[source],
        source_group=AGENCIES_GROUP,
    )


@app.route("/newspapers")
def newspapers_page():
    return render_news_page(
        source_group=NEWSPAPERS_GROUP,
    )


@app.route("/newspapers/found")
def newspapers_found_page():
    return render_news_page(
        mode="found",
        source_group=NEWSPAPERS_GROUP,
    )


@app.route("/newspapers/filter/<path:source>")
def newspapers_filter_source(source):
    return render_news_page(
        source_filters=[source],
        source_group=NEWSPAPERS_GROUP,
    )


@app.route("/article", methods=["GET", "POST"])
def article_page():
    if request.method == "POST" and not csrf_is_valid():
        abort(400)
    url = request.values.get("url", "").strip()
    item = find_news_by_url(url)
    if item is None:
        # Совместимость со старыми тестовыми/JSON-наборами; рабочая SQLite-база
        # находит публикацию индексом и не загружает десятки тысяч записей.
        item = next(
            (
                news
                for news in load_json("all_news.json", [])
                if news.get("url") == url
            ),
            None,
        )
    if item is None:
        abort(404)
    force_refresh = request.method == "POST"
    cached = load_cached_article(url)
    if (
        is_yahoo_source(item.get("source", ""))
        and yahoo_article_is_polluted(cached)
    ):
        # Версия 2026.08.17.6 могла сохранить вместе со статьёй рейтинг Yahoo.
        # Такой кэш не показываем и заменяем чистым текстом при этом открытии.
        cached = None
    embedded_paragraphs = [
        " ".join(str(paragraph).split())
        for paragraph in item.get("article_paragraphs", [])
        if " ".join(str(paragraph).split())
    ]
    if embedded_paragraphs:
        # Некоторые официальные API/ленты сразу содержат полный текст.
        # Он надёжнее повторного разбора декоративной HTML-страницы.
        article = {
            "title": item.get("title", ""),
            "paragraphs": embedded_paragraphs,
            "error": "",
        }
    elif cached is not None and not force_refresh:
        article = cached
    elif item.get("source") == "ТАСС" and item.get("summary"):
        # RSS ТАСС отдаёт официальный анонс сразу и без браузерной проверки.
        # Полную публикацию при необходимости можно открыть по ссылке.
        article = {
            "title": item.get("title", ""),
            "paragraphs": [item["summary"]],
            "error": "",
        }
    else:
        article = extract_article(url, item.get("title", ""))
        if (
            str(item.get("source", "")).startswith("Yahoo! JAPAN ·")
            and item.get("summary")
            and not article.get("paragraphs")
        ):
            # Полный текст Yahoo загружается только при открытии карточки.
            # Если защита сайта его не отдала, сохраняем официальный RSS-анонс.
            article = {
                "title": item.get("title", ""),
                "paragraphs": [item["summary"]],
                "error": "",
            }
        elif (
            item.get("source") == "Минобороны РФ"
            and item.get("summary")
            and not article.get("paragraphs")
        ):
            # Некоторые публикации Минобороны являются видеоматериалами:
            # отдельная страница содержит заголовок и теги, но не текст.
            # В таком случае показываем официальный анонс из карточки ленты.
            article = {
                "title": item.get("title", ""),
                "paragraphs": [item["summary"]],
                "error": "",
            }
    embedded_text_changed = (
        bool(embedded_paragraphs)
        and cached is not None
        and cached.get("paragraphs") != article.get("paragraphs")
    )
    if cached is None or force_refresh or embedded_text_changed:
        saved = save_cached_article(url, article, item.get("source", ""))
        if saved is not None:
            article = saved
        elif force_refresh and cached is not None:
            refresh_error = article.get("error") or (
                "Источник не отдал новый текст — показана сохранённая версия."
            )
            article = dict(cached)
            article["refresh_error"] = refresh_error

    requested_back = request.values.get("back_url", "").strip()
    if requested_back.startswith("/") and not requested_back.startswith("//"):
        back_url = requested_back
    else:
        back_url = (
            request.referrer
            if request.referrer and request.host in request.referrer
            else "/"
        )
    return render_template(
        "article.html",
        article=article,
        item=item,
        back_url=back_url,
        csrf_token=csrf_token(),
    )


@app.route("/api/keywords", methods=["GET", "POST", "DELETE"])
def keywords_api():
    if request.method == "GET":
        return jsonify(keywords=load_keywords())

    if not csrf_is_valid():
        return jsonify(error="Сессия устарела. Обновите страницу."), 400

    payload = request.get_json(silent=True) or {}
    keyword = str(payload.get("keyword", "")).strip()
    if not 2 <= len(keyword) <= 80:
        return jsonify(error="Слово или фраза должны содержать от 2 до 80 символов."), 400

    if request.method == "POST":
        keywords = add_keyword(keyword)
    else:
        keywords = remove_keyword(keyword)
    found = rebuild_found_news()
    return jsonify(keywords=keywords, found_count=len(found))


@app.route("/api/bookmarks", methods=["GET", "POST", "DELETE"])
def bookmarks_api():
    """Переключает сердечко и складывает выбранную новость в «Моё избранное»."""
    user = current_user()
    user_id = user["id"]
    if request.method == "GET":
        return jsonify(
            urls=bookmarked_urls(user_id),
            count=count_bookmarks(user_id),
        )
    if not csrf_is_valid():
        return jsonify(error="Сессия устарела. Обновите страницу."), 400

    payload = request.get_json(silent=True) or {}
    url = str(payload.get("url", "")).strip()
    if request.method == "DELETE":
        remove_bookmark(user_id, url)
    else:
        favorite_folder = ensure_favorites_folder(user_id)
        legacy_urls = payload.get("urls")
        if isinstance(legacy_urls, list):
            # Старый интерфейс хранил сердечки только в localStorage.
            # Один ограниченный пакет переносит их в текущий аккаунт.
            for legacy_url in legacy_urls[:500]:
                item = find_news_by_url(str(legacy_url).strip())
                if item is not None:
                    save_bookmark(user_id, item, favorite_folder["id"])
        else:
            item = find_news_by_url(url)
            if item is None:
                return jsonify(error="Новость не найдена"), 404
            save_bookmark(user_id, item, favorite_folder["id"])
    return jsonify(
        urls=bookmarked_urls(user_id),
        count=count_bookmarks(user_id),
    )


@app.get("/api/news-index")
def news_index_api():
    """Подгружает компактный индекс личных непрочитанных новостей."""
    source_group = str(request.args.get("group", "")).strip().casefold()
    if source_group not in {
        GOVERNMENT_GROUP,
        AGENCIES_GROUP,
        NEWSPAPERS_GROUP,
    }:
        return jsonify(error="Неизвестный раздел источников"), 400
    user = current_user()
    return jsonify(items=list_unread_news_index(
        user["id"], source_group, UNREAD_INDEX_LIMIT,
    ))


@app.post("/api/news-read")
def news_read_api():
    """Синхронизирует личные отметки чтения между устройствами."""
    if not csrf_is_valid():
        return jsonify(error="Сессия устарела. Обновите страницу."), 400
    user = current_user()
    payload = request.get_json(silent=True) or {}
    try:
        if payload.get("all") is True:
            source_group = str(payload.get("source_group", "")).strip().casefold()
            mark_news_group_read(user["id"], source_group)
        else:
            mark_news_read(user["id"], payload.get("url"))
    except ValueError as error:
        return jsonify(error=str(error)), 400
    return jsonify(ok=True)


@app.post("/api/news-read/migrate")
def news_read_migration_api():
    """Один раз переносит старые браузерные отметки в текущий аккаунт."""
    if not csrf_is_valid():
        return jsonify(error="Сессия устарела. Обновите страницу."), 400
    payload = request.get_json(silent=True) or {}
    unread_urls = payload.get("unread_urls")
    if not isinstance(unread_urls, list):
        return jsonify(error="Некорректные отметки чтения"), 400
    try:
        migrated = migrate_legacy_unread(current_user()["id"], unread_urls)
    except ValueError as error:
        return jsonify(error=str(error)), 400
    return jsonify(ok=True, migrated=migrated)


@app.post("/api/source-order")
def source_order_api():
    """Сохраняет личный порядок правой панели для выбранного раздела."""
    user = current_user()
    if not csrf_is_valid():
        return jsonify(error="Сессия устарела. Обновите страницу."), 400
    payload = request.get_json(silent=True) or {}
    source_group = str(payload.get("source_group", "")).strip().casefold()
    if source_group not in {
        GOVERNMENT_GROUP,
        AGENCIES_GROUP,
        NEWSPAPERS_GROUP,
    }:
        return jsonify(error="Неизвестный раздел источников"), 400
    requested = payload.get("sources")
    if not isinstance(requested, list):
        return jsonify(error="Некорректный порядок источников"), 400

    counts = news_source_counts(source_group)
    available = sorted(counts.items(), key=lambda item: (-item[1], item[0]))
    final_order = [name for name, _ in apply_source_order(available, requested)]
    try:
        saved = save_source_order(user["id"], source_group, final_order)
    except ValueError as error:
        return jsonify(error=str(error)), 400
    return jsonify(sources=saved)


@app.post("/api/collection-order")
def collection_order_api():
    """Сохраняет порядок личных подборок после перетаскивания мышкой."""
    user = current_user()
    if not csrf_is_valid():
        return jsonify(error="Сессия устарела. Обновите страницу."), 400
    payload = request.get_json(silent=True) or {}
    requested = payload.get("folder_ids")
    if not isinstance(requested, list):
        return jsonify(error="Некорректный порядок подборок"), 400
    try:
        folders = save_bookmark_folder_order(user["id"], requested)
    except ValueError as error:
        return jsonify(error=str(error)), 400
    return jsonify(folder_ids=[folder["id"] for folder in folders])


@app.post("/api/collection-note-read")
def collection_note_read_api():
    """Сохраняет личную отметку чтения статьи в доступной подборке."""
    user = current_user()
    if not csrf_is_valid():
        return jsonify(error="Сессия устарела. Обновите страницу."), 400
    payload = request.get_json(silent=True) or {}
    is_read = payload.get("is_read")
    if not isinstance(is_read, bool):
        return jsonify(error="Некорректная отметка чтения"), 400
    try:
        saved = set_collection_note_read(
            user["id"], payload.get("folder_id"), payload.get("note_id"), is_read,
        )
    except ValueError as error:
        return jsonify(error=str(error)), 400
    return jsonify(is_read=saved)


if __name__ == "__main__":
    host = _server_host()
    try:
        port = int(environment_value("MONITOR_PORT", "5000"))
    except ValueError:
        port = 5000
    print(f"🌐 http://{host}:{port}")
    if _enabled_setting("FLASK_DEBUG"):
        app.run(debug=True, host=host, port=port)
    else:
        from waitress import serve

        options = {
            "host": host,
            "port": port,
            "threads": 4,
            "ident": "NewsMonitor",
            "expose_tracebacks": False,
            "clear_untrusted_proxy_headers": True,
            "max_request_body_size": 1_000_000,
            "max_request_header_size": 32_768,
            "channel_timeout": 60,
        }
        if _enabled_setting("MONITOR_TRUST_PROXY"):
            options.update(
                trusted_proxy="127.0.0.1",
                trusted_proxy_count=1,
                trusted_proxy_headers={
                    "x-forwarded-for",
                    "x-forwarded-host",
                    "x-forwarded-port",
                    "x-forwarded-proto",
                },
            )
        serve(app, **options)
